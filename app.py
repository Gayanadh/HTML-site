import os, io, json, base64, tempfile, subprocess
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import fitz
from groq import Groq
import anthropic

app = Flask(__name__)
CORS(app)

groq_client   = Groq(api_key=os.environ.get("GROQ_API_KEY")) if os.environ.get("GROQ_API_KEY") else None
claude_client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY")) if os.environ.get("ANTHROPIC_API_KEY") else None

ATS_ORDER = [
    "Professional Summary","Objective","Key Strengths","Skills",
    "Technical Skills","Achievements","Certifications",
    "Professional History","Work Experience","Education",
    "Projects","Awards","Languages","Volunteer"
]

# ─────────────────────────────────────────────────────────────────
#  STEP 1 — Convert file to image(s) for Claude Vision
# ─────────────────────────────────────────────────────────────────
def file_to_base64_images(file_bytes: bytes, filename: str) -> list[dict]:
    """Convert PDF or DOCX to list of base64 PNG images for Claude Vision."""
    images = []

    if filename.lower().endswith(".pdf"):
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        for page in pdf:
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            images.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": base64.standard_b64encode(img_bytes).decode("utf-8")
                }
            })

    elif filename.lower().endswith(".docx"):
        # Convert DOCX → PDF → images via LibreOffice
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, "resume.docx")
            with open(src, "wb") as f:
                f.write(file_bytes)
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, src],
                capture_output=True, timeout=60
            )
            pdf_path = os.path.join(tmpdir, "resume.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    pdf_bytes_tmp = f.read()
                pdf = fitz.open(stream=pdf_bytes_tmp, filetype="pdf")
                for page in pdf:
                    pix = page.get_pixmap(dpi=150)
                    img_bytes = pix.tobytes("png")
                    images.append({
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": base64.standard_b64encode(img_bytes).decode("utf-8")
                        }
                    })
    return images

# ─────────────────────────────────────────────────────────────────
#  STEP 2 — Claude Vision reads the resume layout
# ─────────────────────────────────────────────────────────────────
def claude_vision_read_resume(images: list[dict]) -> dict:
    """
    Send resume images to Claude Vision.
    Returns structured JSON with every section, heading, and content.
    """
    content = images + [{
        "type": "text",
        "text": """Analyze this resume carefully. Extract the COMPLETE structure.

Return ONLY valid JSON in this exact format:
{
  "candidate_name": "...",
  "layout": "single_column | two_column | sidebar",
  "sections": [
    {
      "heading": "exact heading text as it appears",
      "heading_style": "color, bold, underline, font-size if visible",
      "content": ["bullet 1", "bullet 2", "..."],
      "position": "left_column | right_column | full_width"
    }
  ],
  "design": {
    "primary_color": "hex or description",
    "has_photo": true,
    "has_sidebar": true,
    "font_style": "modern | classic | creative"
  }
}

Include EVERY section you see. Be exact with heading names."""
    }]

    response = claude_client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        messages=[{"role": "user", "content": content}]
    )

    raw = response.content[0].text.strip()
    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    return json.loads(raw)

# ─────────────────────────────────────────────────────────────────
#  STEP 3 — Groq classifies + polishes the new sentence
# ─────────────────────────────────────────────────────────────────
def groq_classify_and_polish(resume_structure: dict, new_text: str, template: str) -> dict:
    """
    Given the full resume structure from Claude Vision + the new sentence,
    Groq picks the exact heading and rewrites the sentence.
    Returns { "heading": "...", "polished": "..." }
    """
    existing_headings = [s["heading"] for s in resume_structure.get("sections", [])]

    system = f"""You are an ATS resume expert.
You have the resume structure. Your job:
1. Choose which heading this sentence belongs under.
   Prefer EXISTING headings: {json.dumps(existing_headings)}
   If none fit, pick the best ATS-standard heading from: {json.dumps(ATS_ORDER)}
2. Rewrite the sentence for {template.upper()} style:
   - Google: XYZ formula (Accomplished X as measured by Y by doing Z)
   - Apple: elegant product/user impact, one sentence
   - Amazon: metrics + Leadership Principles, one sentence
   - Default: strong action verb + quantified result

Return ONLY valid JSON:
{{"heading": "<chosen heading>", "polished": "<rewritten sentence>"}}"""

    response = groq_client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": f"New sentence: {new_text}"}
        ],
        max_tokens=200,
        temperature=0.3
    )

    raw = response.choices[0].message.content.strip()
    try:
        result = json.loads(raw)
        return result
    except Exception:
        return {"heading": "Achievements", "polished": new_text}

# ─────────────────────────────────────────────────────────────────
#  STEP 4 — Rebuild DOCX from structure + new content
# ─────────────────────────────────────────────────────────────────
def rebuild_docx(file_bytes: bytes, filename: str,
                 resume_structure: dict, target_heading: str, new_text: str) -> io.BytesIO:
    """
    Open original DOCX, find the correct heading via XML,
    insert the new bullet directly under it.
    Falls back to appending if heading not found.
    """
    doc = Document(io.BytesIO(file_bytes))

    # Find target paragraph index (fuzzy match)
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        if target_heading.lower() in para.text.strip().lower():
            target_idx = i
            break

    if target_idx is not None:
        # Insert bullet right after the heading paragraph
        anchor = doc.paragraphs[target_idx]
        # Find how many bullets already exist under this heading
        # so we insert after them, not before
        insert_after = anchor
        for j in range(target_idx + 1, len(doc.paragraphs)):
            next_para = doc.paragraphs[j]
            # Stop if we hit another heading
            if any(h.lower() in next_para.text.strip().lower()
                   for h in ATS_ORDER if h.lower() != target_heading.lower()):
                break
            if next_para.text.strip():
                insert_after = next_para
    else:
        # Heading doesn't exist — find where it should go in ATS order
        # and create it
        ats_idx = ATS_ORDER.index(target_heading) if target_heading in ATS_ORDER else 99
        insert_before_idx = len(doc.paragraphs) - 1

        for i, para in enumerate(doc.paragraphs):
            for next_heading in ATS_ORDER[ats_idx + 1:]:
                if next_heading.lower() in para.text.strip().lower():
                    insert_before_idx = i
                    break
            if insert_before_idx != len(doc.paragraphs) - 1:
                break

        anchor = doc.paragraphs[insert_before_idx]

        # Create the new heading
        heading_el = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "2E5C9E")
        pPr.append(shd)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "120")
        spacing.set(qn("w:after"), "60")
        pPr.append(spacing)
        heading_el.append(pPr)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "FFFFFF")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        rPr.extend([b, color, sz])
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = target_heading
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run.append(t)
        heading_el.append(run)
        anchor._element.addprevious(heading_el)
        insert_after = doc.paragraphs[doc.paragraphs.index(anchor) - 1
                                      if anchor in doc.paragraphs else 0]

    # Insert the bullet point
    bullet_el = OxmlElement("w:p")
    pPr2 = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "360")
    pPr2.append(ind)
    spacing2 = OxmlElement("w:spacing")
    spacing2.set(qn("w:before"), "60")
    spacing2.set(qn("w:after"), "60")
    pPr2.append(spacing2)
    bullet_el.append(pPr2)
    run2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    color2 = OxmlElement("w:color")
    color2.set(qn("w:val"), "1A3A6B")
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "20")
    rPr2.extend([color2, sz2])
    run2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = f"\u2022  {new_text}"
    t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run2.append(t2)
    bullet_el.append(run2)
    insert_after._element.addnext(bullet_el)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def rebuild_pdf(file_bytes: bytes, resume_structure: dict,
                target_heading: str, new_text: str) -> io.BytesIO:
    """
    For PDF: convert to DOCX, inject, convert back.
    Uses LibreOffice for conversion.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "input.pdf")
        with open(pdf_path, "wb") as f:
            f.write(file_bytes)

        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", tmpdir, pdf_path],
            capture_output=True, timeout=60
        )
        docx_path = os.path.join(tmpdir, "input.docx")
        if not os.path.exists(docx_path):
            return io.BytesIO(file_bytes)

        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        enhanced = rebuild_docx(
            docx_bytes, "input.docx",
            resume_structure, target_heading, new_text
        )

        out_docx = os.path.join(tmpdir, "enhanced.docx")
        with open(out_docx, "wb") as f:
            enhanced.seek(0)
            f.write(enhanced.read())

        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", tmpdir, out_docx],
            capture_output=True, timeout=60
        )

        out_pdf = os.path.join(tmpdir, "enhanced.pdf")
        if os.path.exists(out_pdf):
            with open(out_pdf, "rb") as f:
                result = io.BytesIO(f.read())
                result.seek(0)
                return result

    return io.BytesIO(file_bytes)

# ─────────────────────────────────────────────────────────────────
#  ROUTES
# ─────────────────────────────────────────────────────────────────
@app.route("/polish", methods=["POST"])
def polish():
    """
    Accepts multipart/form-data with:
      - resume (file)
      - text (the sentence to classify)
      - template (google/apple/amazon)
    Returns { polished_text, heading }
    """
    if not claude_client:
        return jsonify({"polished_text": "Error: ANTHROPIC_API_KEY not set.", "heading": "Achievements"}), 500
    if not groq_client:
        return jsonify({"polished_text": "Error: GROQ_API_KEY not set.", "heading": "Achievements"}), 500

    try:
        file     = request.files.get("resume")
        text     = request.form.get("text", "")
        template = request.form.get("template", "google")

        if not file or not text:
            return jsonify({"polished_text": "Upload a resume and enter text.", "heading": "Achievements"}), 400

        file_bytes = file.read()
        filename   = file.filename

        # Step 1: render to images
        images = file_to_base64_images(file_bytes, filename)
        if not images:
            return jsonify({"polished_text": "Could not render file.", "heading": "Achievements"}), 400

        # Step 2: Claude Vision reads structure
        resume_structure = claude_vision_read_resume(images)

        # Step 3: Groq classifies + polishes
        result = groq_classify_and_polish(resume_structure, text, template)

        return jsonify({
            "polished_text":    result.get("polished", text),
            "heading":          result.get("heading", "Achievements"),
            "resume_structure": resume_structure  # pass back to frontend to cache
        })

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"polished_text": f"Error: {str(e)}", "heading": "Achievements"}), 200


@app.route("/upgrade", methods=["POST"])
def upgrade():
    """
    Accepts multipart/form-data with:
      - resume (file)
      - updates (polished sentence)
      - heading (from /polish)
      - format (docx/pdf)
      - resume_structure (JSON string, optional — skips re-read if provided)
    """
    try:
        file             = request.files.get("resume")
        updates          = request.form.get("updates", "").strip()
        heading          = request.form.get("heading", "Achievements").strip()
        requested_format = request.form.get("format", "docx").lower()
        template         = request.form.get("template", "google")
        cached_structure = request.form.get("resume_structure", "")

        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        file_bytes = file.read()
        filename   = file.filename

        # Use cached structure if available, otherwise re-read
        if cached_structure:
            try:
                resume_structure = json.loads(cached_structure)
            except Exception:
                resume_structure = {}
        elif claude_client:
            images = file_to_base64_images(file_bytes, filename)
            resume_structure = claude_vision_read_resume(images) if images else {}
        else:
            resume_structure = {}

        # If no heading was decided yet, classify now
        if not heading and groq_client and resume_structure:
            result  = groq_classify_and_polish(resume_structure, updates, template)
            heading = result.get("heading", "Achievements")
            updates = result.get("polished", updates)

        # Build the output
        if filename.lower().endswith(".pdf"):
            if requested_format == "pdf":
                result_file = rebuild_pdf(file_bytes, resume_structure, heading, updates)
                return send_file(result_file, mimetype="application/pdf",
                                 as_attachment=True, download_name="Optimized_Resume.pdf")
            else:
                result_file = rebuild_pdf(file_bytes, resume_structure, heading, updates)
                return send_file(result_file, mimetype="application/pdf",
                                 as_attachment=True, download_name="Optimized_Resume.pdf")

        elif filename.lower().endswith(".docx"):
            enhanced = rebuild_docx(file_bytes, filename, resume_structure, heading, updates)

            if requested_format == "docx":
                return send_file(
                    enhanced,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    as_attachment=True, download_name="Optimized_Resume.docx"
                )
            # DOCX → PDF
            enhanced.seek(0)
            with tempfile.TemporaryDirectory() as tmpdir:
                src = os.path.join(tmpdir, "resume.docx")
                with open(src, "wb") as f:
                    f.write(enhanced.read())
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf",
                     "--outdir", tmpdir, src],
                    capture_output=True, timeout=60
                )
                pdf_path = os.path.join(tmpdir, "resume.pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        return send_file(io.BytesIO(f.read()), mimetype="application/pdf",
                                         as_attachment=True, download_name="Optimized_Resume.pdf")
            enhanced.seek(0)
            return send_file(enhanced,
                             mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             as_attachment=True, download_name="Optimized_Resume.docx")

        return jsonify({"error": "Unsupported file type"}), 400

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port, debug=False)
